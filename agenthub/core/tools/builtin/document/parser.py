"""
Document Parser Implementation

Provides multi-format document parsing with intelligent content extraction,
metadata extraction, and structured data parsing.
"""

import os
import json
import mimetypes
from typing import Dict, List, Any, Optional
from pathlib import Path
import re

# Document parsing libraries
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    HTML_AVAILABLE = True
except ImportError:
    HTML_AVAILABLE = False

from agenthub.core.tools import tool
from agenthub.core.tools.builtin.base import CachedTool, SecurityValidator


class DocumentParser:
    """Multi-format document parser with intelligent content extraction."""
    
    def __init__(self):
        self.security_validator = SecurityValidator()
        self.supported_formats = {
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.html': self._parse_html,
            '.htm': self._parse_html,
            '.json': self._parse_json,
            '.csv': self._parse_csv,
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.doc': self._parse_docx  # Will try to parse as docx
        }
    
    def parse(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse document with specified options."""
        try:
            # Validate file path
            self.security_validator._validate_file_path(file_path)
            
            # Check if file exists
            if not os.path.exists(file_path):
                return {
                    "success": False,
                    "error": f"File not found: {file_path}",
                    "error_type": "file_not_found"
                }
            
            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            
            # Check if format is supported
            if file_ext not in self.supported_formats:
                return {
                    "success": False,
                    "error": f"Unsupported file format: {file_ext}",
                    "error_type": "unsupported_format",
                    "supported_formats": list(self.supported_formats.keys())
                }
            
            # Parse document
            result = self.supported_formats[file_ext](file_path, options)
            
            # Add common metadata
            result.update({
                "file_path": file_path,
                "file_size": os.path.getsize(file_path),
                "file_extension": file_ext,
                "mime_type": mimetypes.guess_type(file_path)[0]
            })
            
            return result
            
        except Exception as e:
            return {
                "success": False,
                "error": f"Parsing failed: {str(e)}",
                "error_type": "parsing_error"
            }
    
    def _parse_text(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse plain text file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        return {
            "success": True,
            "content": {
                "text": content,
                "word_count": len(content.split()),
                "char_count": len(content),
                "line_count": len(content.splitlines())
            },
            "metadata": {
                "type": "text",
                "encoding": "utf-8"
            }
        }
    
    def _parse_markdown(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse markdown file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Extract headers
        headers = re.findall(r'^#+\s+(.+)$', content, re.MULTILINE)
        
        # Extract links
        links = re.findall(r'\[([^\]]+)\]\(([^)]+)\)', content)
        
        return {
            "success": True,
            "content": {
                "text": content,
                "word_count": len(content.split()),
                "char_count": len(content),
                "line_count": len(content.splitlines()),
                "headers": headers,
                "links": [{"text": link[0], "url": link[1]} for link in links]
            },
            "metadata": {
                "type": "markdown",
                "encoding": "utf-8",
                "header_count": len(headers),
                "link_count": len(links)
            }
        }
    
    def _parse_html(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse HTML file."""
        if not HTML_AVAILABLE:
            return {
                "success": False,
                "error": "BeautifulSoup not available for HTML parsing",
                "error_type": "dependency_missing"
            }
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        soup = BeautifulSoup(content, 'html.parser')
        
        # Extract text content
        text_content = soup.get_text(separator=' ', strip=True)
        
        # Extract metadata
        title = soup.find('title')
        title_text = title.get_text().strip() if title else None
        
        meta_tags = soup.find_all('meta')
        meta_data = {}
        for meta in meta_tags:
            name = meta.get('name') or meta.get('property')
            content_attr = meta.get('content')
            if name and content_attr:
                meta_data[name] = content_attr
        
        # Extract links
        links = [{"text": a.get_text().strip(), "url": a.get('href')} 
                for a in soup.find_all('a', href=True)]
        
        # Extract images
        images = [{"alt": img.get('alt', ''), "src": img.get('src')} 
                 for img in soup.find_all('img', src=True)]
        
        return {
            "success": True,
            "content": {
                "text": text_content,
                "html": content,
                "word_count": len(text_content.split()),
                "char_count": len(text_content),
                "line_count": len(text_content.splitlines()),
                "links": links,
                "images": images
            },
            "metadata": {
                "type": "html",
                "encoding": "utf-8",
                "title": title_text,
                "meta_tags": meta_data,
                "link_count": len(links),
                "image_count": len(images)
            }
        }
    
    def _parse_json(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse JSON file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        return {
            "success": True,
            "content": {
                "data": data,
                "type": type(data).__name__,
                "size": len(str(data))
            },
            "metadata": {
                "type": "json",
                "encoding": "utf-8",
                "is_array": isinstance(data, list),
                "is_object": isinstance(data, dict)
            }
        }
    
    def _parse_csv(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse CSV file."""
        if not PANDAS_AVAILABLE:
            return {
                "success": False,
                "error": "pandas not available for CSV parsing",
                "error_type": "dependency_missing"
            }
        
        try:
            df = pd.read_csv(file_path)
            
            return {
                "success": True,
                "content": {
                    "data": df.to_dict('records'),
                    "columns": df.columns.tolist(),
                    "row_count": len(df),
                    "column_count": len(df.columns)
                },
                "metadata": {
                    "type": "csv",
                    "encoding": "utf-8",
                    "dtypes": df.dtypes.to_dict(),
                    "shape": df.shape
                }
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"CSV parsing failed: {str(e)}",
                "error_type": "csv_parsing_error"
            }
    
    def _parse_pdf(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse PDF file."""
        if not PDF_AVAILABLE:
            return {
                "success": False,
                "error": "PyPDF2 not available for PDF parsing",
                "error_type": "dependency_missing"
            }
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                text_content = ""
                page_count = len(pdf_reader.pages)
                
                for page_num in range(page_count):
                    page = pdf_reader.pages[page_num]
                    text_content += page.extract_text() + "\n"
                
                # Extract metadata
                metadata = pdf_reader.metadata or {}
                
                return {
                    "success": True,
                    "content": {
                        "text": text_content.strip(),
                        "word_count": len(text_content.split()),
                        "char_count": len(text_content),
                        "page_count": page_count
                    },
                    "metadata": {
                        "type": "pdf",
                        "title": metadata.get('/Title', ''),
                        "author": metadata.get('/Author', ''),
                        "subject": metadata.get('/Subject', ''),
                        "creator": metadata.get('/Creator', ''),
                        "producer": metadata.get('/Producer', ''),
                        "creation_date": str(metadata.get('/CreationDate', '')),
                        "modification_date": str(metadata.get('/ModDate', ''))
                    }
                }
        except Exception as e:
            return {
                "success": False,
                "error": f"PDF parsing failed: {str(e)}",
                "error_type": "pdf_parsing_error"
            }
    
    def _parse_docx(self, file_path: str, options: Dict[str, Any]) -> Dict[str, Any]:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            return {
                "success": False,
                "error": "python-docx not available for DOCX parsing",
                "error_type": "dependency_missing"
            }
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text content
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Extract metadata
            core_props = doc.core_properties
            
            return {
                "success": True,
                "content": {
                    "text": text_content.strip(),
                    "word_count": len(text_content.split()),
                    "char_count": len(text_content),
                    "paragraph_count": len(doc.paragraphs),
                    "tables": tables
                },
                "metadata": {
                    "type": "docx",
                    "title": core_props.title or '',
                    "author": core_props.author or '',
                    "subject": core_props.subject or '',
                    "keywords": core_props.keywords or '',
                    "comments": core_props.comments or '',
                    "created": str(core_props.created) if core_props.created else '',
                    "modified": str(core_props.modified) if core_props.modified else '',
                    "table_count": len(tables)
                }
            }
        except Exception as e:
            return {
                "success": False,
                "error": f"DOCX parsing failed: {str(e)}",
                "error_type": "docx_parsing_error"
            }


# Global parser instance
_parser = DocumentParser()


@tool(
    name="document_parse",
    description="Parse documents and extract structured content from multiple formats"
)
def document_parse(
    file_path: str,
    extract_metadata: bool = True,
    extract_tables: bool = True,
    extract_images: bool = False,
    extract_links: bool = False,
    language: str = "auto"
) -> Dict[str, Any]:
    """
    Parse document and extract structured content.
    
    Args:
        file_path: Path to the document file
        extract_metadata: Extract document metadata (author, title, etc.)
        extract_tables: Extract tables as structured data
        extract_images: Extract image references and metadata
        extract_links: Extract hyperlinks and references
        language: Document language for better parsing
    
    Returns:
        dict: Parsed content with text, metadata, tables, and structure
    """
    options = {
        "extract_metadata": extract_metadata,
        "extract_tables": extract_tables,
        "extract_images": extract_images,
        "extract_links": extract_links,
        "language": language
    }
    
    return _parser.parse(file_path, options)
